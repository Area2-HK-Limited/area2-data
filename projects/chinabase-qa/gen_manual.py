#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Chinabase CMS 使用者手冊 DOCX
"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

BASE_DIR = '/home/openclaw/.openclaw/workspace/data/projects/chinabase-qa'
SCREENSHOT_DIR = os.path.join(BASE_DIR, 'manual-screenshots')
OUTPUT_DOCX = os.path.join(BASE_DIR, 'Chinabase_CMS_使用者手冊.docx')

# Chapter to screenshot mapping
CHAPTER_SCREENSHOTS = {
    1: '00_login.png',
    2: 'product.png',
    3: 'category.png',
    4: 'doc_category.png',
    5: 'doc_subcategory.png',
    6: 'news.png',
    7: 'recruitment.png',  # actual file is recruitment.png
    8: 'developing.png',
    9: 'banner.png',
    10: 'certificate.png',
    11: 'subscriber.png',
    12: 'product_inquiry.png',
    13: 'enquiry.png',
    14: 'translate.png',
    15: 'filemanager.png',
}

def set_font(run, size_pt=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = 'Arial'
    # Set East Asian font (微軟正黑體)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    # Check if rFonts already exists
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_paragraph_font(para, size_pt=11):
    """Set font for all runs in a paragraph."""
    for run in para.runs:
        set_font(run, size_pt=size_pt)

def add_styled_heading(doc, text, level):
    """Add a heading with proper font."""
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = 'Arial'
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)
    return para

def add_paragraph_text(doc, text, style='Normal', bold=False, italic=False):
    """Add a paragraph with proper font settings."""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    return para

def add_screenshot(doc, chapter_num):
    """Add screenshot for a chapter if it exists."""
    filename = CHAPTER_SCREENSHOTS.get(chapter_num)
    if not filename:
        return
    img_path = os.path.join(SCREENSHOT_DIR, filename)
    if not os.path.exists(img_path):
        print(f"  ⚠️  Screenshot not found: {img_path}")
        return
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(img_path, width=Cm(14))
        # Add caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cap.add_run(f'圖 {chapter_num}: {filename.replace(".png", "")}')
        run2.font.size = Pt(9)
        run2.italic = True
        run2.font.name = 'Arial'
        print(f"  ✅ Screenshot inserted: {filename}")
    except Exception as e:
        print(f"  ❌ Failed to insert screenshot {filename}: {e}")

def add_table_from_markdown(doc, table_lines):
    """Parse markdown table and add to document."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line or line.startswith('|---') or re.match(r'^\|[-\s|]+\|$', line):
            continue
        # Parse cells
        cells = [c.strip() for c in line.strip('|').split('|')]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Determine column count
    num_cols = max(len(row) for row in rows)
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text
                # Style header row
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'
                else:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'

def parse_and_add_content(doc, lines, chapter_num=None):
    """Parse markdown lines and add to document."""
    i = 0
    in_table = False
    table_lines = []
    step_counter = 0
    in_steps = False
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip screenshot lines (already handled separately)
        if stripped.startswith('![') and 'manual-screenshots' in stripped:
            i += 1
            continue
        
        # Table detection
        if '|' in stripped and stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            # End of table - render it
            add_table_from_markdown(doc, table_lines)
            in_table = False
            table_lines = []
            # Don't increment i, process current line
        
        # Headings
        if stripped.startswith('### '):
            text = stripped[4:]
            add_styled_heading(doc, text, level=3)
            step_counter = 0
            i += 1
            continue
        
        elif stripped.startswith('## '):
            text = stripped[3:]
            add_styled_heading(doc, text, level=2)
            step_counter = 0
            i += 1
            continue
        
        elif stripped.startswith('# '):
            text = stripped[2:]
            add_styled_heading(doc, text, level=1)
            i += 1
            continue
        
        # Numbered list items
        elif re.match(r'^\d+\.\s', stripped):
            step_counter += 1
            text = re.sub(r'^\d+\.\s', '', stripped)
            # Remove markdown bold markers
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # Remove inline code markers
            text = re.sub(r'`([^`]+)`', r'\1', text)
            para = doc.add_paragraph(style='List Number')
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            r = run._r
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '微軟正黑體')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            existing = rPr.find(qn('w:rFonts'))
            if existing is not None:
                rPr.remove(existing)
            rPr.insert(0, rFonts)
            i += 1
            continue
        
        # Bullet list items
        elif stripped.startswith('- '):
            text = stripped[2:]
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            r = run._r
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '微軟正黑體')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            existing = rPr.find(qn('w:rFonts'))
            if existing is not None:
                rPr.remove(existing)
            rPr.insert(0, rFonts)
            i += 1
            continue
        
        # Horizontal rule
        elif stripped == '---':
            doc.add_paragraph()  # spacing
            i += 1
            continue
        
        # Empty line
        elif stripped == '':
            i += 1
            continue
        
        # Regular paragraph text
        else:
            # Clean up markdown formatting
            text = stripped
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # links
            
            if text:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(11)
                run.font.name = 'Arial'
                r = run._r
                rPr = r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), '微軟正黑體')
                rFonts.set(qn('w:ascii'), 'Arial')
                rFonts.set(qn('w:hAnsi'), 'Arial')
                existing = rPr.find(qn('w:rFonts'))
                if existing is not None:
                    rPr.remove(existing)
                rPr.insert(0, rFonts)
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)


def main():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # ---- Title Page ----
    title_para = doc.add_heading('Chinabase CMS 使用者手冊', 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(24)
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)
    
    # Subtitle info
    for info in ['版本：1.0', '日期：2026-03-10', '網站：https://chinabase.hlhk.net/', 'CMS 入口：https://chinabase.hlhk.net/cms/', '適用帳號：user1']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    
    doc.add_page_break()
    
    # ---- Table of Contents ----
    toc_heading = doc.add_heading('目錄', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    toc_entries = [
        '1. 登入系統',
        '2. 產品管理 — 產品',
        '3. 產品管理 — 產品分類',
        '4. 產品管理 — 產品文件分類',
        '5. 產品管理 — 產品文件',
        '6. 內容管理 — 最新消息',
        '7. 內容管理 — 誠聘英才',
        '8. 內容管理 — 發展里程',
        '9. 內容管理 — 橫額圖片',
        '10. 內容管理 — 專利證書',
        '11. 查詢管理 — 電郵訂閱',
        '12. 查詢管理 — 產品查詢',
        '13. 查詢管理 — 客戶反饋',
        '14. 系統工具 — 翻譯管理',
        '15. 系統工具 — 文件管理器',
    ]
    for entry in toc_entries:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(entry)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    
    doc.add_page_break()
    
    # ---- Read and parse manual.md ----
    manual_path = os.path.join(BASE_DIR, 'manual.md')
    with open(manual_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into chapter sections
    # Each chapter starts with "## N. " pattern
    chapter_pattern = re.compile(r'^(## \d+\. .+)$', re.MULTILINE)
    
    # Find all chapter positions
    chapters = list(chapter_pattern.finditer(content))
    
    # Also find the appendix section
    appendix_match = re.search(r'^(## 附錄.+)$', content, re.MULTILINE)
    
    # Process intro (before first chapter)
    # Skip it - we have our own title page
    
    print(f"Found {len(chapters)} chapters")
    
    for idx, chapter_match in enumerate(chapters):
        chapter_num = idx + 1
        
        # Get chapter content
        start = chapter_match.start()
        end = chapters[idx + 1].start() if idx + 1 < len(chapters) else (
            appendix_match.start() if appendix_match else len(content)
        )
        
        chapter_content = content[start:end].strip()
        lines = chapter_content.split('\n')
        
        print(f"\nProcessing Chapter {chapter_num}: {lines[0][:60]}")
        
        # Add chapter heading (first line is "## N. Title")
        chapter_title = lines[0].strip().lstrip('# ').strip()
        add_styled_heading(doc, chapter_title, level=1)
        
        # Process rest of chapter (skip first line - heading already added)
        content_lines = lines[1:]
        
        # Filter out screenshot lines and process content
        filtered_lines = []
        for line in content_lines:
            stripped = line.strip()
            # Skip screenshot image lines
            if stripped.startswith('![') and 'manual-screenshots' in stripped:
                continue
            filtered_lines.append(line)
        
        parse_and_add_content(doc, filtered_lines, chapter_num)
        
        # Add screenshot after chapter content
        print(f"  Adding screenshot for chapter {chapter_num}...")
        add_screenshot(doc, chapter_num)
        
        # Page break between chapters (except last)
        if chapter_num < len(chapters):
            doc.add_page_break()
    
    # ---- Appendix ----
    if appendix_match:
        doc.add_page_break()
        appendix_start = appendix_match.start()
        appendix_content = content[appendix_start:].strip()
        appendix_lines = appendix_content.split('\n')
        
        print(f"\nProcessing Appendix...")
        add_styled_heading(doc, '附錄：常見問題', level=1)
        
        # Process appendix lines (skip the heading line)
        parse_and_add_content(doc, appendix_lines[1:])
    
    # ---- Save ----
    doc.save(OUTPUT_DOCX)
    print(f"\n✅ DOCX saved: {OUTPUT_DOCX}")
    print(f"   Size: {os.path.getsize(OUTPUT_DOCX) / 1024:.1f} KB")


if __name__ == '__main__':
    main()
