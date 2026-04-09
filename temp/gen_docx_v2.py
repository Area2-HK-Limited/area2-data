from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def shade_para(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def heading(text, size=14, color=(26,82,118)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p

def spacer():
    doc.add_paragraph()

def overview_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def day_bar(text, fill='1A5276'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(255, 255, 255)
    shade_para(p, fill)
    return p

def depart(time_str):
    p = doc.add_paragraph()
    run = p.add_run(f'⏰ 建議出門口時間：{time_str}')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(180, 30, 30)
    return p

def time_item(time_str, activity):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{time_str} ')
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor(26, 82, 118)
    r2 = p.add_run(activity)
    r2.font.size = Pt(9.5)
    return p

def sub_item(activity, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(f'• {activity}')
    run.font.size = Pt(9)
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def accom(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(f'🏨 {text}')
    run.font.size = Pt(8.5)
    run.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    return p

# ─── TITLE ───────────────────────────────────────────────────────────────────
heading('峴港 + 會安 家庭旅遊行程表', size=15, color=(26, 82, 118))
heading('越南 6日5夜自由行  ·  2026年4月5日至4月10日', size=9, color=(100,100,100))
spacer()

# ─── OVERVIEW ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
shade_para(p, 'F0F0F5')
r = p.add_run('📋 行程概覽')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(26,82,118)

for line in [
    '出發地：香港  →  目的地：越南峴港 & 會安',
    '行程日期：2026年4月5日（星期日）至 4月10日（星期五）',
    '旅客人數：2位大人 + 1位小朋友（7歲）',
    '家庭成員：Zen、Chloe 及 女兒',
    '住宿①：會安 — Grand Sunrise Palace Hoian（會安升陽皇宮酒店）',
    '住宿②：峴港 — Le Sands Oceanfront Hotel Danang（海風海岸酒店）',
]:
    overview_line(line)
spacer()

# ─── DAY 1 ───────────────────────────────────────────────────────────────────
day_bar('第1天  ·  4月5日（星期日）  ·  抵達日')
p = doc.add_paragraph()
r = p.add_run('⚠️ 今日為抵達日，全日安排較輕鬆')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(80,80,80)
time_item('17:25', '抵達峴港國際機場')
sub_item('機場往會安車程約45分鐘')
sub_item('抵達後直接前往酒店 Check-in：Grand Sunrise Palace Hoian')
accom('住宿：Grand Sunrise Palace Hoian（會安）')
spacer()

# ─── DAY 2 ───────────────────────────────────────────────────────────────────
day_bar('第2天  ·  4月6日（星期一）  ·  會安古鎮一天遊')
depart('08:00  →  建議從酒店出發')
time_item('早上', '占婆島椰林upu Water Palm — 乘坐越南傳統竹籃船、捉螃蟹體驗')
sub_item('秋盆河（Thu Bon River）— 划船穿越棕櫚林，風景優美')
time_item('中午', '會安古鎮（UNESCO聯合國世界文化遺產）')
sub_item('日本廊橋（來遠橋）— 建於18世紀的標誌性古橋')
sub_item('廣肇會館 — 傳統華人會館建築')
sub_item('會安古老大宅 — 百年商人大宅，了解當地歷史')
time_item('下午', '自由活動：選購會安手工燈籠、步行探索古鎮街道風情')
time_item('晚上', '會安夜市 — 色彩繽紛的燈籠夜景、越南街頭小食')
accom('住宿：Grand Sunrise Palace Hoian（會安）')
spacer()

# ─── DAY 3 ───────────────────────────────────────────────────────────────────
day_bar('第3天  ·  4月7日（星期二）  ·  美山聖地 + 峴港市內遊')
depart('07:30  →  建議提早出門，美山聖地需要充足遊覽時間')
time_item('早上', '美山聖地（My Son Sanctuary）— 世界文化遺產，印度教寺廟建築群')
sub_item('建議遊覽時間：1.5-2小時')
sub_item('返回峴港，車程約1小時')
sub_item('抵達後入住：Le Sands Oceanfront Hotel Danang')
time_item('12:30', '午餐：Pizza 4P\'s（峴港分店，適合家庭）')
time_item('下午', '峴港市內觀光 — 粉色教堂、龍橋、金龍橋夜景')
sub_item('龍橋（Dragon Bridge）— 每逢周末晚上9點有噴火噴水表演')
time_item('傍晚', '美溪沙灘（My Khe Beach）— 游泳、沙灘玩樂、欣賞日落')
time_item('19:00', '晚餐：Hai San Moc Quen — 海鮮餐廳，新鮮魚蝦蟹')
accom('住宿：Le Sands Oceanfront Hotel Danang（峴港）')
spacer()

# ─── DAY 4 ───────────────────────────────────────────────────────────────────
day_bar('第4天  ·  4月8日（星期三）  ·  烹飪課程 + 峴港市郊')
depart('08:30  →  烹飪學校集合時間通常為早上，建議預留充裕時間')
time_item('早上', '越南烹飪課程（包午餐）— 學整越南美食，小朋友可以一齊參與，非常適合家庭！')
sub_item('一般包含：市場遊覽（認識食材）→ 烹飪示範 → 親手整3-4道菜 → 一起品嚐午餐')
time_item('下午', '五行山（Marble Mountains）— 探索石灰岩洞穴、寺廟，景色開揚壯觀')
sub_item('山茶半島（Son Tra）— 山茶靈應寺，巨型佛像坐像，俯瞰海岸線全景')
time_item('19:00', '晚餐：GAINN BBQ — 韓式烤肉鐵板燒，適合喜歡吃肉的小朋友')
accom('住宿：Le Sands Oceanfront Hotel Danang（峴港）')
spacer()

# ─── DAY 5 ───────────────────────────────────────────────────────────────────
day_bar('第5天  ·  4月9日（星期四）  ·  巴拿山（Ba Na Hills）一天遊')
depart('07:30  →  建議提早出門，巴拿山幅員廣闊，需要充足時間探索')
time_item('全日', '巴拿山（Ba Na Hills）— 峴港最受歡迎嘅親子景點，必去！')
sub_item('乘坐全世界最長纜車登山（景色壯觀，纜車程約20分鐘）')
sub_item('黃金橋（Golden Bridge）— 標誌性「巨人之手」橋樑，拍照打卡一流')
sub_item('法國村（French Hill）— 歐陸小鎮風情，猶如置身歐洲')
sub_item('夢幻樂園（Fantasy Park）— 室內遊戲機、恐龍展、迷你過山車，小朋友最愛！')
sub_item('Buffet 自助午餐已包含在門票內')
sub_item('山頂氣溫比市區低，建議帶備薄外套')
time_item('傍晚', '返回酒店休息，或享受足部按摩放鬆一下')
time_item('晚餐', 'Le Comptoir 或 Le Petit Chef — 精緻法式/意大利餐，環境適合家庭')
accom('住宿：Le Sands Oceanfront Hotel Danang（峴港）')
spacer()

# ─── DAY 6 ───────────────────────────────────────────────────────────────────
day_bar('第6天  ·  4月10日（星期五）  ·  購物及離開', fill='B41E1E')
p = doc.add_paragraph()
r = p.add_run('⚠️ 今日需要預留充足時間，下午機場手續務必提前完成')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(180,30,30)
depart('09:00  →  建議悠閒吃早餐後出發，11:00前開始購物')
time_item('早上', '漢市場（Han Market）— 峴港最大型傳統市場，購買手信、乾貨、零食')
sub_item('建議選購：咖啡、榴槤餅、腰果、紫薯餅等越南特產')
sub_item('Vincom Center / Go Supermarket — 大型商場，補足所需')
time_item('下午', 'Cafe Hopping — 峴港有很多IG打卡咖啡店，可選1-2間拍拍照休息一下')
sub_item('建議時間：14:00-16:00 之間，之後準備往機場')
time_item('17:50', '抵達峴港國際機場（航班起飛時間 20:50）')
sub_item('辦理登機手續時間：17:50  ｜  登機口關閉：19:20  ｜  起飛：20:50', italic=True, color=(80,80,80))
spacer()

# ─── PAGE BREAK → TIPS ──────────────────────────────────────────────────────
doc.add_page_break()
heading('👨‍👩‍👧 親子旅遊實用建議', size=13, color=(26,82,118))
spacer()

tips = [
    ('⭐ 巴拿山是親子必去！',
     'Fantasy Park室內遊樂場有唔少機動遊戲，建議帶小朋友嘅家長預留起碼4-5個鐘。記得帶外套，山上氣溫比較涼。'),
    ('🐉 龍橋噴火表演（只限周末）',
     '每逢周六及週日晚上9點，龍橋會有免費噴火噴水表演，小朋友睇到好興奮！記得帶毛巾，可能會濕身。'),
    ('🚣 椰林撐艇（Cam Thanh）',
     '建議早上進行，天氣較涼爽。有部分tour包含捉螃蟹及即場整帽，小朋友好鐘意。'),
    ('🏖️ 美溪沙灘最佳到訪時間',
     '傍晚4-5點係最佳時段，可以睇日落又唔會太曬。沙灘闊而且水淺，好適合小朋友游水和玩沙。'),
    ('🏯 會安學整燈籠',
     '古鎮內有唔少店鋪提供15分鐘整燈籠工作坊，可以整完帶走留念，係好好嘅旅遊回憶！'),
    ('🎢 VinWonders Nam Hoi An（如時間允許）',
     '大型主題樂園+水上樂園+野生動物區，如果小朋友係機動遊戲迷，可以考慮加一日前往。'),
    ('🍕 適合小朋友嘅餐廳',
     'Pizza 4P\'s係家庭友善餐廳，大人小朋友都岩食。越南法包（Banh Mi）同鮮蝦肉卷都係唔錯嘅選擇。'),
    ('💡 建議租用私人包車',
     '第2-5日建議租用私人包車連司機（每日約$40-60美元），有小朋友同行會方便好多，唔洗自己揾車。'),
]

for title, body in tips:
    p = doc.add_paragraph()
    r1 = p.add_run(title)
    r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = RGBColor(26,82,118)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(9)
    spacer()

p = doc.add_paragraph()
r = p.add_run('資料來源：Sassymamasg / Friends Travel Vietnam / TravelynnFamily / Tripadvisor  |  由 A2 旅遊策劃整理')
r.italic = True; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(120,120,120)

out = '/home/ubuntu/.openclaw/workspace/data/temp/da_nang_itinerary_v2.docx'
doc.save(out)
print(f'OK: {out}')
