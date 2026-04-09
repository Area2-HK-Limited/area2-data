from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# Styles
def heading(doc, text, size=14, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def day_bar(doc, text):
    """Colored day header bar"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Set paragraph background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A5276')
    pPr.append(shd)
    return p

def line(doc, text, bold=False, italic=False, size=10, indent=0, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def spacer(doc):
    doc.add_paragraph()

# Title
heading(doc, 'Da Nang + Hoi An Family Trip Itinerary', size=16, color=(26, 82, 118))
spacer(doc)

# Overview
p = doc.add_paragraph()
run = p.add_run('Trip Overview')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(26, 82, 118)
p.paragraph_format.space_after = Pt(2)

for item in ['Destination: Da Nang & Hoi An, Vietnam',
              'Dates: April 5 - April 10, 2026 (6 Days 5 Nights)',
              'Travelers: 2 Adults + 1 Child (7 years old)',
              'Family: Zen, Chloe & Daughter']:
    line(doc, item, size=10)
spacer(doc)

# Day 1
day_bar(doc, 'Day 1  |  April 5 (Sunday)  |  Arrival')
line(doc, '5:25 PM: Arrive at Da Nang International Airport', bold=True, size=10, color=(26, 82, 118))
line(doc, '  Taxi to Hoi An (~45 min)', size=10)
line(doc, '  Check-in: Grand Sunrise Palace Hoian', size=10)
line(doc, 'Accommodation: Grand Sunrise Palace Hoian (Hoi An)', italic=True, size=9, color=(80, 80, 80))
spacer(doc)

# Day 2
day_bar(doc, 'Day 2  |  April 6 (Monday)  |  Hoi An Exploration')
line(doc, 'Morning: Cam Thanh Coconut Forest - Bamboo basket boat ride, crab fishing', bold=True, size=10, color=(26, 82, 118))
line(doc, '  Thu Bon River (秋盤河) - Scenic boat ride through water palms', size=10)
line(doc, 'Afternoon: Hoi An Ancient Town - UNESCO-listed old town, lanterns, street life', bold=True, size=10, color=(26, 82, 118))
line(doc, '  Japanese Covered Bridge (來遠橋) - Iconic 18th-century bridge', size=10)
line(doc, '  Quang Chong Association Hall (廣肇會館) - Traditional assembly hall', size=10)
line(doc, '  Ancient Houses (會安古老大宅) - Historic merchant residences', size=10)
line(doc, 'Evening: Hoi An Night Market - Lanterns, street food, souvenirs', bold=True, size=10, color=(26, 82, 118))
line(doc, 'Accommodation: Grand Sunrise Palace Hoian (Hoi An)', italic=True, size=9, color=(80, 80, 80))
spacer(doc)

# Day 3
day_bar(doc, 'Day 3  |  April 7 (Tuesday)  |  My Son + Da Nang City')
line(doc, 'Morning: My Son Sanctuary - UNESCO Hindu temple ruins (1.5-2 hrs)', bold=True, size=10, color=(26, 82, 118))
line(doc, '  Drive back to Da Nang (~1 hr) → Check-in: Le Sands Oceanfront Hotel Danang', size=10)
line(doc, "12:30 PM: Lunch: Pizza 4P's - Hoang Van Thu (family pizzeria)", bold=True, size=10, color=(26, 82, 118))
line(doc, 'Afternoon: Da Nang City - Pink Cathedral, Dragon Bridge views, Foot massage', bold=True, size=10, color=(26, 82, 118))
line(doc, 'Late PM: My Khe Beach (美溪沙灘) - Swimming, sand play, sunset', bold=True, size=10, color=(26, 82, 118))
line(doc, '7:00 PM: Dinner: Hai San Moc Quen - Fresh seafood restaurant', bold=True, size=10, color=(26, 82, 118))
line(doc, 'Accommodation: Le Sands Oceanfront Hotel Danang (Da Nang)', italic=True, size=9, color=(80, 80, 80))
spacer(doc)

# Day 4
day_bar(doc, 'Day 4  |  April 8 (Wednesday)  |  Cooking Class + Sightseeing')
line(doc, 'Morning: Cooking Class (with lunch) - Learn Vietnamese cuisine, hands-on fun for kids!', bold=True, size=10, color=(26, 82, 118))
line(doc, 'Afternoon: Marble Mountains (五行山) - Cave exploration, pagodas, scenic views', bold=True, size=10, color=(26, 82, 118))
line(doc, '  Linh Ung Pagoda, Son Tra (山茶半島) - Giant Buddha statue, coastal views', size=10)
line(doc, '7:00 PM: Dinner: GAINN BBQ - Korean-Hanoi style grilled meat', bold=True, size=10, color=(26, 82, 118))
line(doc, 'Accommodation: Le Sands Oceanfront Hotel Danang (Da Nang)', italic=True, size=9, color=(80, 80, 80))
spacer(doc)

# Day 5
day_bar(doc, 'Day 5  |  April 9 (Thursday)  |  Ba Na Hills Day Trip')
line(doc, 'All Day: Ba Na Hills (巴拿山) - Full day recommended', bold=True, size=10, color=(26, 82, 118))
line(doc, "  World's longest cable car ride (spectacular mountain views)", size=10)
line(doc, '  The Golden Bridge "God\'s Hand" - iconic landmark, great for photos', size=10)
line(doc, '  French Hill, Deboud winery, Alpine Garden, Ancient Castle', size=10)
line(doc, '  Fantasy Park (室內遊樂城) - Arcade, dinosaurs, mini coaster', size=10)
line(doc, '  Buffet Lunch included', size=10)
line(doc, 'Evening: Hotel nearby - Relaxation / Foot massage', bold=True, size=10, color=(26, 82, 118))
line(doc, 'Dinner: Le Comptoir OR Le Petit Chef (fine dining, family-friendly)', bold=True, size=10, color=(26, 82, 118))
line(doc, 'Accommodation: Le Sands Oceanfront Hotel Danang (Da Nang)', italic=True, size=9, color=(80, 80, 80))
spacer(doc)

# Day 6
p = doc.add_paragraph()
run = p.add_run('Day 6  |  April 10 (Friday)  |  Shopping & Departure')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(255, 255, 255)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'B41E1E')
pPr.append(shd)

line(doc, 'Morning: Han Market (漢市場) - Souvenirs, dried seafood, local snacks', bold=True, size=10, color=(26, 82, 118))
line(doc, '  Vincom / Go Supermarket - Additional souvenir shopping', size=10)
line(doc, '  Instagram-worthy cafes - cafe hopping, decor spots', size=10)
p2 = doc.add_paragraph()
r2 = p2.add_run('5:50 PM: Arrive Da Nang Airport (8:50 PM departure)')
r2.bold = True
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(180, 30, 30)
line(doc, 'Check-in: 5:50 PM  |  Gate closes ~8:20 PM  |  Flight: 8:50 PM', size=9)
spacer(doc)

# Page break
doc.add_page_break()

# Tips page
heading(doc, 'Kid-Friendly Tips & Recommended Additions', size=14, color=(26, 82, 118))
spacer(doc)

tips = [
    ('⭐ Ba Na Hills is a MUST with kids',
     "The world's longest cable car is an adventure itself. Fantasy Park has arcade games, dinosaur displays, and a mini roller coaster — perfect for a 7-year-old. Dress warm at the top (cooler than city). Allow full day."),
    ('🐉 Dragon Bridge Fire Show (Weekends only)',
     "Every Saturday & Sunday at 9 PM — the dragon breathes fire and water, kids absolutely love it. Free from Han River bank. Bring a towel, you might get wet!"),
    ('🚣 Coconut Basket Boat (Cam Thanh)',
     'Bamboo boat spinning through coconut groves — thrilling for kids. Many tours include crab catching and hat-making. Great photo opportunity.'),
    ('🏖️ My Khe & An Bang Beach',
     'Wide, clean, shallow-water beach. Perfect for kids to swim and play sand. Best visited late afternoon for sunset.'),
    ('🏯 Lantern Making (Hoi An)',
     'Many shops offer 15-min lantern-making classes. Kids can make and take home their own silk lantern. A lovely keepsake from the trip.'),
    ('🎢 VinWonders Nam Hoi An',
     'Huge amusement park + waterpark + animal safari. Full-day with wave pools, rides (some 1.3m+ height limits), tiger/elephant shows. Worth considering if kids love theme parks.'),
    ('🍕 Food Tips for Kids',
     "Pizza 4P's (Day 3 lunch) is very kid-friendly. Banh mi and fresh spring rolls are usually big hits. Vietnamese iced coffee is sweet with lots of milk — kids love it too!"),
    ('💡 General Tips',
     'Rent a private car with driver for Day 2-5 (~$40-60/day). Easier with a 7-year-old. Most attractions offer free entry for kids under 1m or under 6. Pack swim gear, light jacket for Ba Na Hills, mosquito repellent.'),
]

for title, body in tips:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(26, 82, 118)
    line(doc, body, size=9.5)
    spacer(doc)

p = doc.add_paragraph()
run = p.add_run('Research: Sassymamasg, Friends Travel Vietnam, TravelynnFamily, Tripadvisor | Compiled by A2 Travel Planner')
run.italic = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(100, 100, 100)

out_path = '/home/ubuntu/.openclaw/workspace/data/temp/da_nang_itinerary.docx'
doc.save(out_path)
print(f'OK: {out_path}')
