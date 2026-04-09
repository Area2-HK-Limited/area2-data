#!/usr/bin/env python3
"""Generate Word document with spine MRI images"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

IMG_DIR = '/home/ubuntu/.openclaw/workspace/data/temp/spine_selected'
OUT_DOCX = '/home/ubuntu/.openclaw/workspace/data/temp/spine_pdf/YUAN_YU_CHAN_MRI_Spine_Selected.docx'

images = sorted([f for f in os.listdir(IMG_DIR) if f.endswith('.jpg')],
                 key=lambda x: int(x.split('_')[0]))
print(f"Total images: {len(images)}")

doc = Document()

# Set page to A4
section = doc.sections[0]
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(1.5)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('脊椎 MRI 影像報告')
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(26, 74, 122)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run('Patient: YUAN YU CHAN  |  MR Whole Spine  |  2026-03-31')
sub_run.font.size = Pt(10)
sub_run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()  # spacer

total_pages = (len(images) + 3) // 4
IMG_W = Inches(3.3)

for i, img_name in enumerate(images):
    if i > 0 and i % 4 == 0:
        doc.add_page_break()
        # Mini header on new pages
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(f'— Page {i//4 + 1}/{total_pages} —')
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(150, 150, 150)
        doc.add_paragraph()

    img_path = os.path.join(IMG_DIR, img_name)
    label = re.sub(r'^\d+_', '', img_name).replace('.jpg', '')

    # Image centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(img_path, width=IMG_W)
    except Exception as e:
        print(f"  Warning [{img_name}]: {e}")

    # Label
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(f'{i+1}. {label}')
    lr.font.size = Pt(7)
    lr.font.color.rgb = RGBColor(100, 100, 100)

    # Spacer between rows
    if i % 2 == 1:
        doc.add_paragraph()

    print(f"  [{i+1:02d}] {img_name}")

doc.save(OUT_DOCX)
print(f"\n✅ Word saved: {OUT_DOCX}")